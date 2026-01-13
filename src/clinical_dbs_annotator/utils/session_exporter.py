"""
Session data exporter for Clinical DBS Annotator.

This module provides functionality to export session data to Excel format.
"""

import os
from datetime import datetime
from typing import Optional

import pandas as pd
from PyQt5.QtWidgets import QMessageBox, QWidget
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

class SessionExporter:
    """
    Handles exporting session data to various formats.
    
    This class provides methods to export the collected session data
    to Excel, with plans for future formats like PDF and Word.
    """
    
    def __init__(self, session_data):
        """
        Initialize the session exporter.
        
        Args:
            session_data: The SessionData instance containing collected data
        """
        self.session_data = session_data
     
    
    def _read_session_data(self) -> Optional[pd.DataFrame]:
        """
        Read session data from the TSV file.
        
        Returns:
            DataFrame with session data or None if error
        """
        try:
            if hasattr(self.session_data, 'file_path') and self.session_data.file_path:
                return pd.read_csv(self.session_data.file_path, sep='\t')
            return None
        except Exception:
            return None
    
    def _create_summary_sheet(self, writer, df: pd.DataFrame) -> None:
        """
        Create a summary sheet with statistics.
        
        Args:
            writer: Excel writer object
            df: Session data DataFrame
        """
        summary_data = []
        
        # Basic statistics
        summary_data.append(['Total Records', len(df)])
        summary_data.append(['Unique Scales', df['scale_name'].nunique() if 'scale_name' in df.columns else 0])
        
        # Stimulation parameter ranges
        if 'left_amplitude' in df.columns:
            summary_data.append(['Left Amplitude Range', f"{df['left_amplitude'].min():.1f} - {df['left_amplitude'].max():.1f} mA"])
        if 'right_amplitude' in df.columns:
            summary_data.append(['Right Amplitude Range', f"{df['right_amplitude'].min():.1f} - {df['right_amplitude'].max():.1f} mA"])
        
        # Scale statistics
        if 'scale_name' in df.columns and 'scale_value' in df.columns:
            for scale_name in df['scale_name'].unique():
                scale_data = df[df['scale_name'] == scale_name]['scale_value']
                summary_data.append([
                    f'{scale_name} Range',
                    f"{scale_data.min():.1f} - {scale_data.max():.1f}"
                ])
        
        summary_df = pd.DataFrame(summary_data, columns=['Metric', 'Value'])
        summary_df.to_excel(writer, sheet_name='Summary', index=False)
    
    def _auto_adjust_columns(self, worksheet, df: pd.DataFrame) -> None:
        """
        Auto-adjust column widths in Excel worksheet.
        
        Args:
            worksheet: Excel worksheet object
            df: DataFrame with data
        """
        from openpyxl.utils import get_column_letter
        
        for idx, col in enumerate(df.columns, 1):
            # Find the maximum length in the column
            max_length = max(
                len(str(col)),
                df[col].astype(str).str.len().max()
            )
            # Adjust width (with some padding)
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[get_column_letter(idx)].width = adjusted_width


    def export_to_excel(self, parent: Optional[QWidget] = None) -> bool:   
        """
        Export session data to Excel format.
        
        Args:
            parent: Parent widget for dialog display
            
        Returns:
            True if export was successful, False otherwise
        """
        try:
            # Get the current session data
            if not self.session_data.is_file_open():
                QMessageBox.warning(
                    parent, 
                    "No Session Data", 
                    "No session file is currently open. Please start a session first."
                )
                return False
            
            # Read the TSV data
            df = self._read_session_data()
            if df is None or df.empty:
                QMessageBox.warning(
                    parent,
                    "No Data to Export",
                    "No session data has been recorded yet."
                )
                return False
            
            # Generate default filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"dbs_session_report_{timestamp}.xlsx"
            
            # Get save location
            from PyQt5.QtWidgets import QFileDialog
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Session Report",
                default_filename,
                "Excel Files (*.xlsx);;All Files (*)"
            )
            
            if not file_path:
                return False  # User cancelled
            
            # Ensure .xlsx extension
            if not file_path.endswith('.xlsx'):
                file_path += '.xlsx'
            
            # Create Excel file with formatting
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                # Main session data
                df.to_excel(writer, sheet_name='Session Data', index=False)
                
                # Summary sheet
                self._create_summary_sheet(writer, df)
                
                # Get access to the workbook for formatting
                workbook = writer.book
                worksheet = writer.sheets['Session Data']
                
                # Auto-adjust column widths
                self._auto_adjust_columns(worksheet, df)
            
            # Show success message
            QMessageBox.information(
                parent,
                "Export Successful",
                f"Session report exported successfully to:\n{file_path}"
            )
            
            return True
            
        except Exception as e:
            QMessageBox.critical(
                parent,
                "Export Error",
                f"Failed to export session data:\n{str(e)}"
            )
            return False

    def export_to_pdf(self, parent: Optional[QWidget] = None) -> bool:
        """
        Export session data to PDF format.
        
        Args:
            parent: Parent widget for dialog display
            
        Returns:
            True if export was successful, False otherwise
        """
        try:
            # Get the current session data
            if not self.session_data.is_file_open():
                QMessageBox.warning(
                    parent, 
                    "No Session Data", 
                    "No session file is currently open. Please start a session first."
                )
                return False
            
            # Read the TSV data
            df = self._read_session_data()
            if df is None or df.empty:
                QMessageBox.warning(
                    parent,
                    "No Data to Export",
                    "No session data has been recorded yet."
                )
                return False
            
            # Generate default filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"dbs_session_report_{timestamp}.pdf"
            
            # Get save location
            from PyQt5.QtWidgets import QFileDialog
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Session Report",
                default_filename,
                "PDF Files (*.pdf);;All Files (*)"
            )
            
            if not file_path:
                return False  # User cancelled
            
            # Ensure .pdf extension
            if not file_path.endswith('.pdf'):
                file_path += '.pdf'
            
            # Create PDF document
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title_style = styles['Title']
            title = Paragraph("Clinical DBS Session Report", title_style)
            story.append(title)
            
            # Add metadata
            metadata_style = styles['Normal']
            metadata_text = f"""
            Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}<br/>
            Total Records: {len(df)}<br/>
            Unique Scales: {df["scale_name"].nunique() if "scale_name" in df.columns else 0}
            """
            story.append(Paragraph(metadata_text, metadata_style))
            story.append(Paragraph("<br/>", metadata_style))
            
            # Add summary section
            heading_style = styles['Heading1']
            story.append(Paragraph("Session Summary", heading_style))
            
            # Create summary table
            summary_data = [
                ['Metric', 'Value'],
                ['Total Records', str(len(df))],
                ['Unique Scales', str(df["scale_name"].nunique() if "scale_name" in df.columns else 0)],
            ]
            
            # Stimulation parameter ranges
            if 'left_amplitude' in df.columns:
                summary_data.append([
                    'Left Amplitude Range',
                    f"{df['left_amplitude'].min():.1f} - {df['left_amplitude'].max():.1f} mA"
                ])
            if 'right_amplitude' in df.columns:
                summary_data.append([
                    'Right Amplitude Range',
                    f"{df['right_amplitude'].min():.1f} - {df['right_amplitude'].max():.1f} mA"
                ])
            
            # Create summary table
            summary_table = Table(summary_data)
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(summary_table)
            story.append(Paragraph("<br/>", metadata_style))
            
            # Add data table section
            story.append(Paragraph("Session Data", heading_style))
            
            # Prepare data for table
            table_data = [list(df.columns)]
            for _, row in df.iterrows():
                table_data.append([str(value) if pd.notna(value) else '' for value in row])
            
            # Create main data table
            data_table = Table(table_data)
            
            # Style the table
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ])
            
            # Alternate row colors for better readability
            for i in range(1, len(table_data)):
                if i % 2 == 0:
                    table_style.add('BACKGROUND', (0, i), (-1, i), colors.lightgrey)
            
            data_table.setStyle(table_style)
            story.append(data_table)
            
            # Add notes section if notes exist
            if 'notes' in df.columns:
                story.append(Paragraph("<br/>", metadata_style))
                story.append(Paragraph("Session Notes", heading_style))
                
                unique_notes = df['notes'].dropna().unique()
                for note in unique_notes:
                    if note.strip():
                        note_text = f"• {note}"
                        story.append(Paragraph(note_text, metadata_style))
            
            # Build PDF
            doc.build(story)
            
            # Show success message
            QMessageBox.information(
                parent,
                "Export Successful",
                f"Session report exported successfully to:\n{file_path}"
            )
            
            return True
            
        except Exception as e:
            QMessageBox.critical(
                parent,
                "Export Error",
                f"Failed to export session data to PDF:\n{str(e)}"
            )
            return False
    
    def export_to_word(self, parent: Optional[QWidget] = None) -> bool:
        """
        Export session data to Word format.
        
        Args:
            parent: Parent widget for dialog display
            
        Returns:
            True if export was successful, False otherwise
        """
        try:
            # Get the current session data
            if not self.session_data.is_file_open():
                QMessageBox.warning(
                    parent, 
                    "No Session Data", 
                    "No session file is currently open. Please start a session first."
                )
                return False
            
            # Read the TSV data
            df = self._read_session_data()
            if df is None or df.empty:
                QMessageBox.warning(
                    parent,
                    "No Data to Export",
                    "No session data has been recorded yet."
                )
                return False
            
            # Generate default filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"dbs_session_report_{timestamp}.docx"
            
            # Get save location
            from PyQt5.QtWidgets import QFileDialog
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Session Report",
                default_filename,
                "Word Files (*.docx);;All Files (*)"
            )
            
            if not file_path:
                return False  # User cancelled
            
            # Ensure .docx extension
            if not file_path.endswith('.docx'):
                file_path += '.docx'
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Clinical DBS Session Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total Records: {len(df)}')
            doc.add_paragraph('')  # Empty line
            
            # Add summary section
            doc.add_heading('Session Summary', level=1)
            
            # Basic statistics
            summary_data = [
                f'Total Records: {len(df)}',
                f'Unique Scales: {df["scale_name"].nunique() if "scale_name" in df.columns else 0}',
            ]
            
            # Stimulation parameter ranges
            if 'left_amplitude' in df.columns:
                summary_data.append(
                    f'Left Amplitude Range: {df["left_amplitude"].min():.1f} - {df["left_amplitude"].max():.1f} mA'
                )
            if 'right_amplitude' in df.columns:
                summary_data.append(
                    f'Right Amplitude Range: {df["right_amplitude"].min():.1f} - {df["right_amplitude"].max():.1f} mA'
                )
            
            for item in summary_data:
                doc.add_paragraph(item, style='List Bullet')
            
            doc.add_paragraph('')  # Empty line
            
            # Add data table
            doc.add_heading('Session Data', level=1)
            
            # Convert DataFrame to table
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                hdr_cells[i].text = str(col_name).replace('_', ' ').title()
                # Make header bold
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows
            for i, (_, row) in enumerate(df.iterrows()):
                row_cells = table.rows[i + 1].cells
                for j, value in enumerate(row):
                    row_cells[j].text = str(value) if pd.notna(value) else ''
            
            doc.add_paragraph('')  # Empty line
            
            # Add notes section if notes exist
            if 'notes' in df.columns:
                doc.add_heading('Session Notes', level=1)
                unique_notes = df['notes'].dropna().unique()
                for note in unique_notes:
                    if note.strip():
                        doc.add_paragraph(f'• {note}', style='List Bullet')
            
            # Save the document
            doc.save(file_path)
            
            # Show success message
            QMessageBox.information(
                parent,
                "Export Successful",
                f"Session report exported successfully to:\n{file_path}"
            )
            
            return True
            
        except Exception as e:
            QMessageBox.critical(
                parent,
                "Export Error",
                f"Failed to export session data to Word:\n{str(e)}"
            )
            return False
