import sys
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QTableWidget, 
                             QTableWidgetItem, QFileDialog, QMessageBox, QLabel)
from PyQt5.QtCore import Qt
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExcelPDFGenerator(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()
        
    def initUI(self):
        """Inizializza l'interfaccia utente"""
        self.setWindowTitle('Generatore Excel/PDF')
        self.setGeometry(100, 100, 800, 600)
        
        # Widget centrale
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # Label informativa
        info_label = QLabel('Inserisci i dati nella tabella sottostante:')
        layout.addWidget(info_label)
        
        # Tabella per inserire i dati
        self.table = QTableWidget(10, 5)  # 10 righe, 5 colonne
        self.table.setHorizontalHeaderLabels(['Nome', 'Cognome', 'Età', 'Email', 'Città'])
        layout.addWidget(self.table)
        
        # Pulsanti di controllo
        button_layout = QHBoxLayout()
        
        add_row_btn = QPushButton('Aggiungi Riga')
        add_row_btn.clicked.connect(self.add_row)
        button_layout.addWidget(add_row_btn)
        
        clear_btn = QPushButton('Pulisci Tabella')
        clear_btn.clicked.connect(self.clear_table)
        button_layout.addWidget(clear_btn)
        
        sample_btn = QPushButton('Dati di Esempio')
        sample_btn.clicked.connect(self.load_sample_data)
        button_layout.addWidget(sample_btn)
        
        layout.addLayout(button_layout)
        
        # Pulsanti di esportazione
        export_layout = QHBoxLayout()
        
        excel_btn = QPushButton('Esporta in Excel')
        excel_btn.clicked.connect(self.export_to_excel)
        excel_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        export_layout.addWidget(excel_btn)
        
        word_btn = QPushButton('Esporta in Word')
        word_btn.clicked.connect(self.export_to_word)
        word_btn.setStyleSheet("background-color: #2B579A; color: white; font-weight: bold;")
        export_layout.addWidget(word_btn)
        
        pdf_btn = QPushButton('Esporta in PDF')
        pdf_btn.clicked.connect(self.export_to_pdf)
        pdf_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold;")
        export_layout.addWidget(pdf_btn)
        
        layout.addLayout(export_layout)
        
        # Pulsante per esportare tutti i formati
        all_export_layout = QHBoxLayout()
        
        all_btn = QPushButton('Esporta Tutti i Formati (Excel + Word + PDF)')
        all_btn.clicked.connect(self.export_all)
        all_btn.setStyleSheet("background-color: #FF9800; color: white; font-weight: bold; padding: 10px;")
        all_export_layout.addWidget(all_btn)
        
        layout.addLayout(all_export_layout)
        
    def add_row(self):
        """Aggiunge una nuova riga alla tabella"""
        row_count = self.table.rowCount()
        self.table.insertRow(row_count)
        
    def clear_table(self):
        """Pulisce tutti i dati dalla tabella"""
        self.table.clearContents()
        
    def load_sample_data(self):
        """Carica dati di esempio nella tabella"""
        sample_data = [
            ['Mario', 'Rossi', '35', 'mario.rossi@email.it', 'Roma'],
            ['Laura', 'Bianchi', '28', 'laura.bianchi@email.it', 'Milano'],
            ['Giuseppe', 'Verdi', '42', 'giuseppe.verdi@email.it', 'Napoli'],
            ['Anna', 'Ferrari', '31', 'anna.ferrari@email.it', 'Torino'],
            ['Paolo', 'Esposito', '39', 'paolo.esposito@email.it', 'Bologna']
        ]
        
        for row_idx, row_data in enumerate(sample_data):
            for col_idx, value in enumerate(row_data):
                self.table.setItem(row_idx, col_idx, QTableWidgetItem(value))
                
    def get_table_data(self):
        """Estrae i dati dalla tabella in formato DataFrame"""
        headers = []
        for col in range(self.table.columnCount()):
            headers.append(self.table.horizontalHeaderItem(col).text())
        
        data = []
        for row in range(self.table.rowCount()):
            row_data = []
            has_data = False
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                if item and item.text():
                    has_data = True
                    row_data.append(item.text())
                else:
                    row_data.append('')
            if has_data:
                data.append(row_data)
        
        return pd.DataFrame(data, columns=headers) if data else None
    
    def export_to_excel(self):
        """Esporta i dati in formato Excel"""
        df = self.get_table_data()
        if df is None or df.empty:
            QMessageBox.warning(self, 'Attenzione', 'Nessun dato da esportare!')
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'Salva File Excel', '', 'Excel Files (*.xlsx)')
        
        if file_path:
            try:
                # Crea un writer Excel con formattazione
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name='Dati', index=False)
                    
                    # Ottieni il workbook e worksheet per la formattazione
                    workbook = writer.book
                    worksheet = writer.sheets['Dati']
                    
                    # Formatta l'intestazione
                    from openpyxl.styles import Font, PatternFill, Alignment
                    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                    header_font = Font(bold=True, color='FFFFFF')
                    
                    for cell in worksheet[1]:
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = Alignment(horizontal='center')
                    
                    # Adatta la larghezza delle colonne
                    for column in worksheet.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            if cell.value:
                                max_length = max(max_length, len(str(cell.value)))
                        worksheet.column_dimensions[column_letter].width = max_length + 2
                
                QMessageBox.information(self, 'Successo', f'File Excel salvato: {file_path}')
            except Exception as e:
                QMessageBox.critical(self, 'Errore', f'Errore durante il salvataggio: {str(e)}')
    
    def export_to_word(self):
        """Esporta i dati in formato Word"""
        df = self.get_table_data()
        if df is None or df.empty:
            QMessageBox.warning(self, 'Attenzione', 'Nessun dato da esportare!')
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'Salva File Word', '', 'Word Files (*.docx)')
        
        if file_path:
            try:
                doc = Document()
                
                # Aggiungi titolo
                title = doc.add_heading('Report Dati', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Aggiungi spazio
                doc.add_paragraph()
                
                # Crea tabella (righe = dati + 1 intestazione, colonne = numero colonne)
                table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
                table.style = 'Light Grid Accent 1'
                
                # Aggiungi intestazioni
                header_cells = table.rows[0].cells
                for idx, column_name in enumerate(df.columns):
                    cell = header_cells[idx]
                    cell.text = str(column_name)
                    # Formatta l'intestazione
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    # Colora lo sfondo dell'intestazione
                    shading_elm = cell._element.get_or_add_tcPr()
                    shading = shading_elm.get_or_add_shd()
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '4472C4')
                
                # Aggiungi dati
                for row_idx, row_data in df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row_data):
                        cell = row_cells[col_idx]
                        cell.text = str(value)
                        # Formatta il testo
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                
                # Centra la tabella
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Salva il documento
                doc.save(file_path)
                
                QMessageBox.information(self, 'Successo', f'File Word salvato: {file_path}')
            except Exception as e:
                QMessageBox.critical(self, 'Errore', f'Errore durante il salvataggio: {str(e)}')
    
    def export_to_pdf(self):
        """Esporta i dati in formato PDF"""
        df = self.get_table_data()
        if df is None or df.empty:
            QMessageBox.warning(self, 'Attenzione', 'Nessun dato da esportare!')
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'Salva File PDF', '', 'PDF Files (*.pdf)')
        
        if file_path:
            try:
                doc = SimpleDocTemplate(file_path, pagesize=A4)
                elements = []
                
                # Titolo
                styles = getSampleStyleSheet()
                title = Paragraph("<b>Report Dati</b>", styles['Title'])
                elements.append(title)
                elements.append(Paragraph("<br/><br/>", styles['Normal']))
                
                # Prepara i dati per la tabella
                table_data = [df.columns.tolist()] + df.values.tolist()
                
                # Crea la tabella
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
                ]))
                
                elements.append(pdf_table)
                doc.build(elements)
                
                QMessageBox.information(self, 'Successo', f'File PDF salvato: {file_path}')
            except Exception as e:
                QMessageBox.critical(self, 'Errore', f'Errore durante il salvataggio: {str(e)}')
    
    def export_all(self):
        """Esporta in tutti i formati: Excel, Word e PDF"""
        df = self.get_table_data()
        if df is None or df.empty:
            QMessageBox.warning(self, 'Attenzione', 'Nessun dato da esportare!')
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, 'Salva File (senza estensione)', '', 'All Files (*)')
        
        if file_path:
            # Rimuovi l'estensione se presente
            base_path = file_path.rsplit('.', 1)[0]
            
            success_messages = []
            error_messages = []
            
            # Esporta Excel
            excel_path = base_path + '.xlsx'
            try:
                with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name='Dati', index=False)
                    workbook = writer.book
                    worksheet = writer.sheets['Dati']
                    
                    from openpyxl.styles import Font, PatternFill, Alignment
                    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                    header_font = Font(bold=True, color='FFFFFF')
                    
                    for cell in worksheet[1]:
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = Alignment(horizontal='center')
                    
                    for column in worksheet.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            if cell.value:
                                max_length = max(max_length, len(str(cell.value)))
                        worksheet.column_dimensions[column_letter].width = max_length + 2
                success_messages.append(f'✓ Excel: {excel_path}')
            except Exception as e:
                error_messages.append(f'✗ Excel: {str(e)}')
            
            # Esporta Word
            word_path = base_path + '.docx'
            try:
                doc = Document()
                title = doc.add_heading('Report Dati', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                
                table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
                table.style = 'Light Grid Accent 1'
                
                header_cells = table.rows[0].cells
                for idx, column_name in enumerate(df.columns):
                    cell = header_cells[idx]
                    cell.text = str(column_name)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    shading_elm = cell._element.get_or_add_tcPr()
                    shading = shading_elm.get_or_add_shd()
                    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '4472C4')
                
                for row_idx, row_data in df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row_data):
                        cell = row_cells[col_idx]
                        cell.text = str(value)
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.save(word_path)
                success_messages.append(f'✓ Word: {word_path}')
            except Exception as e:
                error_messages.append(f'✗ Word: {str(e)}')
            
            # Esporta PDF
            pdf_path = base_path + '.pdf'
            try:
                doc = SimpleDocTemplate(pdf_path, pagesize=A4)
                elements = []
                
                styles = getSampleStyleSheet()
                title = Paragraph("<b>Report Dati</b>", styles['Title'])
                elements.append(title)
                elements.append(Paragraph("<br/><br/>", styles['Normal']))
                
                table_data = [df.columns.tolist()] + df.values.tolist()
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
                ]))
                
                elements.append(pdf_table)
                doc.build(elements)
                success_messages.append(f'✓ PDF: {pdf_path}')
            except Exception as e:
                error_messages.append(f'✗ PDF: {str(e)}')
            
            # Mostra risultati
            message = ''
            if success_messages:
                message += 'File salvati con successo:\n\n' + '\n'.join(success_messages)
            if error_messages:
                if message:
                    message += '\n\n'
                message += 'Errori:\n\n' + '\n'.join(error_messages)
            
            if error_messages:
                QMessageBox.warning(self, 'Completato con errori', message)
            else:
                QMessageBox.information(self, 'Successo', message)

def main():
    app = QApplication(sys.argv)
    window = ExcelPDFGenerator()
    window.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()